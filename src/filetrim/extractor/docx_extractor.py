from __future__ import annotations

from pathlib import Path

from docx import Document

from filetrim.extractor.base import BaseExtractor
from filetrim.file_types import FileType, detect_file_type
from filetrim.models import ExtractedContent


class DocxExtractor(BaseExtractor):
    supported_types = (FileType.DOCX,)

    def extract(self, path: str | Path) -> ExtractedContent:
        source_path = Path(path)
        file_type = detect_file_type(source_path)

        if not self.can_extract(file_type):
            raise ValueError(f"DocxExtractor does not support: {file_type}")

        document = Document(source_path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        text = "\n".join(paragraph for paragraph in paragraphs if paragraph)

        return ExtractedContent(
            source_path=source_path,
            file_type=file_type,
            suffix=source_path.suffix.lower(),
            text=text,
            metadata={
                "paragraph_count": str(len(document.paragraphs)),
                "source_name": source_path.name,
            },
        )
