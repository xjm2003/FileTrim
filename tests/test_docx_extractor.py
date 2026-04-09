from pathlib import Path

from docx import Document

from filetrim.extractor.docx_extractor import DocxExtractor
from filetrim.file_types import FileType


def test_docx_extractor_reads_docx_file(tmp_path: Path) -> None:
    sample = tmp_path / "sample.docx"

    document = Document()
    document.add_paragraph("Resume")
    document.add_paragraph("Education")
    document.save(sample)

    result = DocxExtractor().extract(sample)

    assert result.file_type is FileType.DOCX
    assert result.suffix == ".docx"
    assert result.text == "Resume\nEducation"
    assert result.metadata["paragraph_count"] == "2"
